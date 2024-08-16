import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX

from PyQt6.QtCore import *  # temp
from PyQt6.QtGui import *
from PyQt6.QtWidgets import *

from encodingcontroller import hex_to_tuple_rgb

# from the svg recognised colour names (https://www.w3.org/TR/SVG11/types.html#ColorKeywords)
colours = ((0, 0, 0, WD_COLOR_INDEX.BLACK),
            (255, 0, 0, WD_COLOR_INDEX.BLUE),
            (0, 255, 127, WD_COLOR_INDEX.BRIGHT_GREEN),
            (0, 0, 139, WD_COLOR_INDEX.DARK_BLUE),
            (139, 0, 0, WD_COLOR_INDEX.DARK_RED),
            (218, 165, 32, WD_COLOR_INDEX.DARK_YELLOW),
            (192, 192, 192, WD_COLOR_INDEX.GRAY_25),
            (128, 128, 128, WD_COLOR_INDEX.GRAY_50),
            (0, 128, 0, WD_COLOR_INDEX.GREEN),
            (255, 192, 203, WD_COLOR_INDEX.PINK),
            (255, 0, 0, WD_COLOR_INDEX.RED),
            (0, 128, 128, WD_COLOR_INDEX.TEAL),
            (64, 228, 208, WD_COLOR_INDEX.TURQUOISE),
            (238, 130, 238, WD_COLOR_INDEX.VIOLET),
            (255, 255, 255, WD_COLOR_INDEX.WHITE),
            (255, 255, 0, WD_COLOR_INDEX.YELLOW),
            )

def find_nearest_colour(queried_colour):
    # we dont need the alpha value (4th index) that is given from QColor
    new_colour = (queried_colour[0], queried_colour[1], queried_colour[2])
    
    return min(colours, key = lambda subject: sum((colour - queried) ** 2 for colour, queried in zip(subject, new_colour)))

def apply_formatting(character_format : QTextCharFormat, temp_run):
    if character_format.font().bold():
        temp_run.bold = True
    if character_format.font().underline():
        temp_run.underline = True
    if character_format.font().italic():
        temp_run.italic = True
    if temp_font_family := character_format.font().family():
        temp_run.font.name = temp_font_family
    if temp_font_size := character_format.font().pointSize():
        temp_run.font.size = Pt(temp_font_size)
    if temp_font_colour := character_format.foreground().color().getRgb():
        temp_run.font.color.rgb = RGBColor(temp_font_colour[0], temp_font_colour[1], temp_font_colour[2])
    if temp_font_highlight_colour := character_format.background().color().getRgb():
        temp_run.highlight_color = find_nearest_colour(temp_font_highlight_colour)

# def parse_docx_file_to_html(docx_file_path: str):
#     with open(docx_file_path, "rb") as docx_file:
#         result = mammoth.convert_to_html(docx_file)
#         html = result.value
#         messages = result.messages
#         # print(messages)
#         print(html)
#         return html
def dump(obj):
    for attr in dir(obj):
        print("obj.%s = %r" % (attr, getattr(obj, attr)))
        
def parse_docx_file_to_list(docx_file_path : str):
    docx_file = Document(docx_file_path)
    file_formatting_list = []
    
    def docx_colour_to_qcolor(colour):
        return QColor.fromRgb(*colour)
    
    def qcolor_from_colours(queried_colour : WD_COLOR_INDEX):
        for colour in colours:
            if colour[3] == queried_colour:
                return QColor.fromRgb(colour[0], colour[1], colour[2])
        return QColor().fromRgb(0,0,0,0)
    
    for paragraph in docx_file.paragraphs:
        for temp_run in paragraph.runs:
            temp_format = QTextCharFormat()
            temp_font = temp_run.font
            
            temp_format.setFontFamily(temp_font.name)
            temp_format.setFontUnderline(temp_font.underline or False)
            temp_format.setFontItalic(temp_font.italic or False)
            temp_format.setFontWeight(QFont.Weight.Bold if temp_font.bold else QFont.Weight.Normal)
            temp_format.setFontPointSize(temp_font.size)
            temp_format.setBackground(qcolor_from_colours(str(temp_font.highlight_color)))
            temp_format.setForeground(docx_colour_to_qcolor(hex_to_tuple_rgb(str(temp_font.color.rgb))))
            
            formatted_list = (temp_run.text, temp_format)
            file_formatting_list.append(formatted_list)
            #print(str(temp_font.highlight_color))
    return file_formatting_list

            
def parse_editor_document_to_docx(editor_document: QTextDocument):
    doc = Document()
    paragraph = doc.add_paragraph()

    cursor = QTextCursor(editor_document)
    cursor.movePosition(QTextCursor.MoveOperation.Start)

    while not cursor.atEnd():
        cursor.clearSelection()
        cursor.movePosition(QTextCursor.MoveOperation.NextCharacter, QTextCursor.MoveMode.KeepAnchor)
        cursor_format = cursor.charFormat()
        if cursor_format.isImageFormat():
            url = QUrl(cursor.charFormat().toImageFormat().name())
            image = editor_document.resource(QTextDocument.ResourceType.ImageResource, url)
            image = QImage(image)
            buffer = QBuffer()
            buffer.open(QBuffer.OpenModeFlag.ReadWrite)
            image.save(buffer, "PNG")

            temp_run = paragraph.add_run()
            temp_run.add_picture(io.BytesIO(buffer.data()))
        elif cursor_format.isCharFormat():
            selected_text = cursor.selectedText() # this returns newlines as unicode
            if (u"\u2029" in selected_text): # convert the page break separator (unicode) back to newline character
                selected_text = "\n" # replace with newline character
            temp_run = paragraph.add_run(selected_text)
            apply_formatting(cursor_format, temp_run)

    return doc


# def get_text_formatting(editor_document: QTextDocument):
#     cursor = QTextCursor(editor_document)

#     formatted_document = []
#     current_character_index = 0

#     while not cursor.atEnd():
#         cursor.setPosition(current_character_index)
#         cursor.movePosition(
#             QTextCursor.MoveOperation.NextCharacter, QTextCursor.MoveMode.KeepAnchor
#         )


#         # if len(cursor.selectedText()) == 1:
#             # print("single character", cursor.selection())
#             # cursor_selection = cursor.charFormat()
#             # if cursor_selection.isCharFormat():
#             #     print(cursor_selection)
#             #     char_format = cursor.charFormat()
#             #     formatted_document.append((
#             #         char_format,
#             #         {
#             #             "bold": (
#             #                 True
#             #                 if char_format.fontWeight() == QFont.Weight.Bold
#             #                 else False
#             #             ),
#             #             "italic": char_format.fontItalic(),
#             #             "underline": char_format.fontUnderline(),
#             #         },
#             #     ))
#             # elif cursor.IsImageFormat():
#             #     print(cursor_selection.toImageFormat())
#             #     print("image format")
#             # else:
#             #     print("unknown format")
#         current_character_index += 1

#     return formatted_document


# def convert_to_docx(editor_document: QTextDocument):
#     text_format = get_text_formatting(editor_document)

#     print(text_format)

#     document = Document()
#     paragraph = document.add_paragraph()
#     for run_text, font_format in text_format:
#         print(run_text, font_format)
#         run = paragraph.add_run(run_text)
#         if text_format["bold"]:
#             run.bold = True
#         if text_format["italic"]:
#             run.italic = True
#         if text_format["underline"]:
#             run.underline = True

#         # if text_format[i]["bold"]:
#         #     document.add_paragraph(editor_document.toPlainText()[i], style="Bold")
#         # elif text_format[i]["italic"]:
#         #     document.add_paragraph(editor_document.toPlainText()[i], style="Italic")
#         # elif text_format[i]["underline"]:
#         #     document.add_paragraph(editor_document.toPlainText()[i], style="Underline")
#         # else:
#         #     document.add_paragraph(editor_document.toPlainText()[i])

#     return document
